# -*- coding: utf-8 -*-
"""
build_docx.py — 한국과학예술융합학회 규격 조판본(.docx) 생성

HWP에서 열어 최종 미세조정하기 위한 조판본이다. 규정 4.1/4.3 반영:
  용지 160mm x 232mm, 표지 1단(1~2면) / 본문 2단(3면부터, 단 너비 75mm,
  간격 10mm, 구분선 0.12mm), 글꼴 표제=중고딕 / 본문=신명조,
  크기(표지 제목 13, 국문이름 12, 소속 11, 영문제목 12, 목차/Abstract/본문
  표제 10, 본문 9.5, 표/그림 제목 9, 표 내용 7~8, 참고문헌 9),
  장평 95, 자간 -5, 줄간격 150(본문) / 130(표지).
표 캡션은 표 위 좌측 [Table-00], 그림 캡션은 그림 아래 가운데 <Figure-01>.
"""

import os
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SRC = os.path.join(BASE, "paper", "draft_ksaforum_en.md")
OUT = os.path.join(BASE, "paper", "KSAF_manuscript_typeset.docx")
FIG_DESIGN = os.path.join(BASE, "paper", "figures", "figure_design.png")
FIG_REG = os.path.join(BASE, "paper", "figures", "figure1_register_gap.png")

GOTHIC = "중고딕"
MYUNGJO = "신명조"


def set_cols(section, num, space_mm=10.0, sep=True):
    sectPr = section._sectPr
    cols = sectPr.xpath("./w:cols")
    cols = cols[0] if cols else OxmlElement("w:cols")
    if cols.getparent() is None:
        sectPr.append(cols)
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), str(int(space_mm * 56.7)))  # mm → twips
    cols.set(qn("w:sep"), "1" if sep else "0")
    cols.set(qn("w:equalWidth"), "1")


def style_run(run, font, size, bold=False, spacing=-5, scale=95):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for a in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(a), font)
    # 자간(1/20 pt 단위) + 장평
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:val"), str(int(spacing)))
    rPr.append(sp)
    w = OxmlElement("w:w")
    w.set(qn("w:val"), str(scale))
    rPr.append(w)


def para(doc, text, font=MYUNGJO, size=9.5, bold=False, align=None,
         line=150, indent=True, space_before=0, space_after=0):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = line / 100.0
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if indent and font == MYUNGJO:
        pf.first_line_indent = Pt(9.5)      # 두 칸 들여쓰기(규정 4.3)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    style_run(r, font, size, bold)
    return p


def parse_md(path):
    """마크다운 초안에서 본문(Abstract~Endnote)을 구조화해 추출."""
    txt = open(path, encoding="utf-8").read()
    body = txt[txt.find("## Abstract"):txt.find("## Submission checklist")]
    blocks = []
    cur = None
    for raw in body.splitlines():
        line = raw.rstrip()
        if line.startswith("## "):
            cur = {"type": "h1", "text": line[3:].strip(), "paras": [],
                   "tables": []}
            blocks.append(cur)
        elif line.startswith("### "):
            blocks.append({"type": "h2", "text": line[4:].strip()})
            cur = blocks[-1]
        elif line.startswith("|"):
            blocks.append({"type": "row", "cells":
                           [c.strip() for c in line.strip("|").split("|")]})
        elif line.startswith("[Table-"):
            blocks.append({"type": "tcap", "text": line})
        elif line.strip():
            if blocks and blocks[-1].get("type") == "p":
                blocks[-1]["text"] += " " + line.strip()
            else:
                blocks.append({"type": "p", "text": line.strip()})
        else:
            if blocks and blocks[-1].get("type") == "p":
                blocks.append({"type": "gap"})
    return blocks


def clean(s):
    s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)
    s = re.sub(r"\[(Table-\d+)\]", r"[\1]", s)
    return s.replace("—", "-").strip()


def main():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Mm(160), Mm(232)
    for m in ("left_margin", "right_margin"):
        setattr(sec, m, Mm(20))
    sec.top_margin, sec.bottom_margin = Mm(20), Mm(20)
    set_cols(sec, 1)

    C = WD_ALIGN_PARAGRAPH.CENTER
    # ---------- 표지 1면 ----------
    for _ in range(3):
        doc.add_paragraph()
    para(doc, "접지 합성 페르소나 조건화가 소형 한국어 언어모델 "
              "미세조정에 미치는 효과", GOTHIC, 13, True, C, 130, False)
    para(doc, "- 사전등록 기반 무페르소나 및 임의 페르소나 대비 통제 실험 -",
         GOTHIC, 13, True, C, 130, False)
    doc.add_paragraph()
    para(doc, "김 환*, **  ·  조 근 태***", GOTHIC, 12, True, C, 130, False)
    para(doc, "* 서울사이버대학교 AI융합대학", GOTHIC, 11, False, C, 130, False)
    para(doc, "** 성균관대학교 일반대학원 기술경영학과", GOTHIC, 11, False, C,
         130, False)
    para(doc, "*** 성균관대학교 시스템경영공학과", GOTHIC, 11, False, C, 130,
         False)
    doc.add_paragraph()
    doc.add_paragraph()
    para(doc, "Effects of Grounded Synthetic Persona Conditioning on "
              "Fine-Tuning a Small Korean Language Model", GOTHIC, 12, True,
         C, 130, False)
    para(doc, "- A Preregistered Controlled Comparison Against No-Persona "
              "and Ad-Hoc Persona Data -", GOTHIC, 12, True, C, 130, False)
    doc.add_paragraph()
    para(doc, "Howard Kim*, **  ·  Keun Tae Cho***", GOTHIC, 11, True, C,
         130, False)
    para(doc, "* College of AI Convergence, Seoul Cyber University, Seoul "
              "01133, South Korea (e-mail: howardkim@iscu.ac.kr)", GOTHIC, 11,
         False, C, 130, False)
    para(doc, "** Department of Management of Technology, Graduate School, "
              "Sungkyunkwan University, Suwon 16419, South Korea "
              "(e-mail: howardkim@skku.edu)", GOTHIC, 11, False, C, 130, False)
    para(doc, "*** Department of Systems Management Engineering, "
              "Sungkyunkwan University, Suwon 16419, South Korea", GOTHIC, 11,
         False, C, 130, False)
    for _ in range(6):
        doc.add_paragraph()
    para(doc, "*** Corresponding Author : Keun Tae Cho "
              "(e-mail: ktcho@skku.edu)", MYUNGJO, 9, False, None, 130, False)
    doc.add_paragraph()
    para(doc, "THE KOREAN SOCIETY OF SCIENCE & ART", GOTHIC, 11, True, C,
         130, False)
    para(doc, "한국과학예술융합학회", GOTHIC, 11, True, C, 130, False)

    # ---------- 표지 2면(빈 페이지) ----------
    doc.add_page_break()
    para(doc, "", MYUNGJO, 9.5, False, None, 130, False)

    # ---------- 3면부터 본문(2단) ----------
    new = doc.add_section(WD_SECTION.NEW_PAGE)
    new.page_width, new.page_height = Mm(160), Mm(232)
    new.left_margin = new.right_margin = Mm(20)
    new.top_margin = new.bottom_margin = Mm(20)
    set_cols(new, 2, 10.0, True)

    # ---------- 목차 (규정 4.3: 본문 2단의 첫 항목, 4.4: 2단계까지) ----------
    para(doc, "Contents", GOTHIC, 10, True, None, 150, False, space_after=4)
    CONTENTS = [
        (0, "Abstract"),
        (0, "I. Introduction"),
        (1, "1.1 Background and Purpose"),
        (1, "1.2 Method and Scope"),
        (0, "II. Theoretical Background"),
        (1, "2.1 Persona-Driven Synthetic Data"),
        (1, "2.2 Grounded Personas and the Verification Gap"),
        (0, "III. Method"),
        (1, "3.1 Design and Preregistration"),
        (1, "3.2 Data Generation and Fine-Tuning"),
        (1, "3.3 Evaluation"),
        (0, "IV. Results"),
        (1, "4.1 Manipulation Check and Confirmatory Tests"),
        (1, "4.2 Exploratory Analyses"),
        (0, "V. Conclusion"),
        (0, "Reference"),
        (0, "Endnote"),
    ]
    for lvl, txt in CONTENTS:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing = 1.5
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        if lvl:
            pf.left_indent = Pt(10)
        style_run(p.add_run(txt), GOTHIC, 9.5, bold=True)
    doc.add_paragraph()

    blocks = parse_md(SRC)
    pending_rows = []

    def flush_table():
        nonlocal pending_rows
        if not pending_rows:
            return
        rows = [r for r in pending_rows
                if not all(set(c) <= set("-: ") for c in r)]
        if rows:
            t = doc.add_table(rows=len(rows), cols=len(rows[0]))
            t.style = "Table Grid"
            for i, r in enumerate(rows):
                for j, cell in enumerate(r):
                    if j >= len(t.columns):
                        continue
                    c = t.cell(i, j)
                    c.text = ""
                    p = c.paragraphs[0]
                    p.paragraph_format.line_spacing = 1.5
                    run = p.add_run(clean(cell))
                    style_run(run, MYUNGJO, 8, bold=(i == 0))
        pending_rows = []

    for b in blocks:
        t = b.get("type")
        if t == "row":
            pending_rows.append(b["cells"])
            continue
        flush_table()
        if t == "h1":
            txt = clean(b["text"])
            para(doc, txt, GOTHIC, 10, True, None, 150, False,
                 space_before=10, space_after=3)
        elif t == "h2":
            para(doc, clean(b["text"]), GOTHIC, 10, True, None, 150, False,
                 space_before=6, space_after=2)
        elif t == "tcap":
            para(doc, clean(b["text"]), MYUNGJO, 9, False,
                 WD_ALIGN_PARAGRAPH.LEFT, 150, False, space_before=6)
        elif t == "p":
            txt = clean(b["text"])
            if txt.startswith("**Key Words**") or txt.startswith("Key Words"):
                para(doc, "Key Words", GOTHIC, 10, True, None, 150, False,
                     space_before=6)
                para(doc, txt.split("**:", 1)[-1].lstrip(": ").strip(),
                     MYUNGJO, 9.5, False, None, 150, False)
            elif txt.startswith("[") and "]" in txt[:6]:
                para(doc, txt, MYUNGJO, 9, False, None, 150, False)
            else:
                para(doc, txt, MYUNGJO, 9.5, False,
                     WD_ALIGN_PARAGRAPH.JUSTIFY, 150, True)
    flush_table()

    # ---------- 그림 2종 ----------
    # 조판 시 Figure-01은 3.1절, Figure-02는 4.2절 근처로 이동한다.
    for path, cap in (
        (FIG_DESIGN, "<Figure-01> Study design: three conditions, identical "
                     "generation and tuning, three preregistered evaluations"),
        (FIG_REG, "<Figure-02> Monotonic decrease of the register age gap by "
                  "conditioning strength (exploratory; filled markers "
                  "indicate 95% CI excluding zero)"),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Mm(72))
        para(doc, cap, MYUNGJO, 9, False, WD_ALIGN_PARAGRAPH.CENTER, 150,
             False)

    doc.save(OUT)
    print("저장:", OUT)


if __name__ == "__main__":
    main()
